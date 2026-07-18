from __future__ import annotations

import logging
from dataclasses import dataclass, field
from pathlib import Path
from typing import Callable

import chardet
from docx import Document as DocxDocument
from docx.table import Table
from docx.text.paragraph import Paragraph
from pypdf import PdfReader

from app.services.ocr_service import OCRDetectionDecision, OCRService
from app.utils.text import clean_text

logger = logging.getLogger(__name__)


@dataclass(slots=True)
class LoadedSection:
    """解析后的逻辑片段。

    对外统一成“section”结构，是为了后续支持：
    - 按页切分 PDF；
    - 按标题切分 Markdown / Docx；
    - 更细粒度的结构化抽取。
    """

    text: str
    metadata: dict[str, object] = field(default_factory=dict)


@dataclass(slots=True)
class LoadedDocument:
    """加载器输出的统一文档结构。"""

    filename: str
    file_type: str
    parser_name: str
    sections: list[LoadedSection]

    @property
    def full_text(self) -> str:
        """返回完整正文，便于后续需要文档级摘要或统计时复用。"""

        return '\n\n'.join(section.text for section in self.sections if section.text.strip())


def _detect_text_encoding(file_bytes: bytes) -> str:
    """用 chardet 对文本文件做编码探测。"""

    detected = chardet.detect(file_bytes)
    encoding = detected.get('encoding') or 'utf-8'
    return encoding


def _load_text_file(file_path: Path) -> LoadedDocument:
    """加载纯文本文件。"""

    file_bytes = file_path.read_bytes()
    encoding = _detect_text_encoding(file_bytes)
    text = file_bytes.decode(encoding, errors='replace')
    loaded_document = LoadedDocument(
        filename=file_path.name,
        file_type='txt',
        parser_name='plain_text_loader',
        sections=[LoadedSection(text=clean_text(text), metadata={'section_type': 'full_text', 'section_index': 0})],
    )
    logger.info(
        '[PARSER] selected: file=%s parser=NATIVE_%s sections=%s file_type=%s',
        file_path.name,
        loaded_document.parser_name.upper(),
        len(loaded_document.sections),
        loaded_document.file_type,
    )
    return loaded_document


def _split_markdown_sections(text: str) -> list[LoadedSection]:
    """按 Markdown 标题做最基础分段。

    这样能为后续“按标题回看 chunk 来源”预留结构信息。
    """

    lines = text.splitlines()
    sections: list[LoadedSection] = []
    current_title = 'Introduction'
    current_lines: list[str] = []
    section_index = 0

    def flush_current_section() -> None:
        nonlocal section_index
        content = clean_text('\n'.join(current_lines))
        if not content:
            return
        sections.append(
            LoadedSection(
                text=content,
                metadata={
                    'section_type': 'markdown_heading',
                    'section_title': current_title,
                    'section_index': section_index,
                },
            )
        )
        section_index += 1

    for line in lines:
        if line.lstrip().startswith('#'):
            flush_current_section()
            current_title = line.lstrip('#').strip() or 'Untitled'
            current_lines = [line]
            continue
        current_lines.append(line)

    flush_current_section()
    return sections or [LoadedSection(text=clean_text(text), metadata={'section_type': 'full_text', 'section_index': 0})]


def _load_markdown_file(file_path: Path) -> LoadedDocument:
    """加载 Markdown 文件。"""

    file_bytes = file_path.read_bytes()
    encoding = _detect_text_encoding(file_bytes)
    text = file_bytes.decode(encoding, errors='replace')
    loaded_document = LoadedDocument(
        filename=file_path.name,
        file_type='md',
        parser_name='markdown_loader',
        sections=_split_markdown_sections(text),
    )
    logger.info(
        '[PARSER] selected: file=%s parser=NATIVE_%s sections=%s file_type=%s',
        file_path.name,
        loaded_document.parser_name.upper(),
        len(loaded_document.sections),
        loaded_document.file_type,
    )
    return loaded_document


def _build_ocr_loaded_document(
    *,
    file_path: Path,
    file_type: str,
    markdown_text: str,
    decision: OCRDetectionDecision,
) -> LoadedDocument:
    """把 OCR 返回的 Markdown 统一转换成当前系统的 `LoadedDocument`。"""

    # 将 OCR 出来的整篇 Markdown 喂给原生 Markdown 解析器进行按标题切分
    sections = _split_markdown_sections(markdown_text)
    decision_metadata = decision.to_metadata()
    for section in sections:
        sec_type = section.metadata.get('section_type')
        # 确保 OCR 返回的即使没有标题，也走 markdown_heading 以进入半结构化切分
        if sec_type == 'full_text':
            sec_type = 'markdown_heading'
            
        section.metadata = {
            **section.metadata,
            **decision_metadata,
            'section_type': sec_type or 'markdown_heading',
            'ocr_used': True,
        }

    return LoadedDocument(
        filename=file_path.name,
        file_type=file_type,
        parser_name='ocr_markdown_loader',
        sections=sections,
    )


def _try_load_with_ocr(file_path: Path, *, file_type: str) -> LoadedDocument | None:
    """按启发式判断尝试走 OCR 解析。

    这个函数的设计目标是“增强而不破坏”：
    - 判断命中时优先用 OCR；
    - OCR 不可用或失败时，调用方仍可继续走原生解析；
    - `.doc` 这种原生解析能力较弱的格式可以直接依赖 OCR。
    """

    ocr_service = OCRService()
    decision = ocr_service.analyze_document(file_path)
    if not decision.should_use_ocr:
        logger.info(
            '[OCR] skipped: file=%s file_type=%s enabled=%s reasons=%s text_chars=%s images=%s tables=%s empty_blocks=%s total_blocks=%s',
            file_path.name,
            file_type,
            ocr_service.can_use_ocr(),
            decision.reasons,
            decision.extracted_text_chars,
            decision.image_count,
            decision.table_count,
            decision.empty_block_count,
            decision.total_block_count,
        )
        print("\n==================================================")
        print(f"👉 文档解析路线: 【原生解析】 (跳过 OCR)")
        print(f"👉 文件名称: {file_path.name}")
        print("==================================================\n")
        logger.info('[PARSER] OCR routing rejected. Proceeding to use NATIVE parser for %s.', file_path.name)
        return None

    logger.info(
        '[OCR] selected: file=%s file_type=%s reasons=%s text_chars=%s images=%s tables=%s',
        file_path.name,
        file_type,
        decision.reasons,
        decision.extracted_text_chars,
        decision.image_count,
        decision.table_count,
    )
    
    print("\n==================================================")
    print(f"👉 文档解析路线: 【OCR 智能识别】")
    print(f"👉 文件名称: {file_path.name}")
    print(f"👉 触发原因: {decision.reasons}")
    print("==================================================\n")

    markdown_text = ocr_service.parse_to_markdown(file_path)
    loaded_document = _build_ocr_loaded_document(
        file_path=file_path,
        file_type=file_type,
        markdown_text=markdown_text,
        decision=decision,
    )
    logger.info(
        '[PARSER] selected: file=%s parser=OCR_%s sections=%s file_type=%s',
        file_path.name,
        loaded_document.parser_name.upper(),
        len(loaded_document.sections),
        loaded_document.file_type,
    )
    return loaded_document


def _load_pdf_file(file_path: Path) -> LoadedDocument:
    """按页加载 PDF 文件。"""

    try:
        ocr_loaded_document = _try_load_with_ocr(file_path, file_type='pdf')
        if ocr_loaded_document is not None:
            return ocr_loaded_document
    except Exception as exc:  # noqa: BLE001
        # OCR 是增强解析链路，不应影响原生文本可提取的 PDF 入库。
        logger.warning('[OCR] PDF parsing failed, fallback to native parser: file=%s error=%s', file_path.name, exc)
        print("\n==================================================")
        print(f"👉 文档解析路线: 【原生解析】 (OCR 失败自动降级)")
        print(f"👉 文件名称: {file_path.name}")
        print(f"👉 失败原因: {exc}")
        print("==================================================\n")

    reader = PdfReader(str(file_path))
    sections: list[LoadedSection] = []
    for page_index, page in enumerate(reader.pages):
        page_text = clean_text(page.extract_text() or '')
        if not page_text:
            continue
        sections.append(
            LoadedSection(
                text=page_text,
                metadata={
                    'section_type': 'pdf_page',
                    'section_index': page_index,
                    'page_number': page_index + 1,
                },
            )
        )

    loaded_document = LoadedDocument(
        filename=file_path.name,
        file_type='pdf',
        parser_name='pypdf_loader',
        sections=sections,
    )
    logger.info(
        '[PARSER] selected: file=%s parser=NATIVE_%s sections=%s file_type=%s',
        file_path.name,
        loaded_document.parser_name.upper(),
        len(loaded_document.sections),
        loaded_document.file_type,
    )
    return loaded_document


def _load_docx_file(file_path: Path) -> LoadedDocument:
    """按标题块加载 Docx 文件。"""

    try:
        ocr_loaded_document = _try_load_with_ocr(file_path, file_type='docx')
        if ocr_loaded_document is not None:
            return ocr_loaded_document
    except Exception as exc:  # noqa: BLE001
        # 对 Docx 来说，OCR 也是增强路线；失败后回退原生解析。
        logger.warning('[OCR] Docx parsing failed, fallback to native parser: file=%s error=%s', file_path.name, exc)
        print("\n==================================================")
        print(f"👉 文档解析路线: 【原生解析】 (OCR 失败自动降级)")
        print(f"👉 文件名称: {file_path.name}")
        print(f"👉 失败原因: {exc}")
        print("==================================================\n")

    document = DocxDocument(str(file_path))
    sections: list[LoadedSection] = []
    current_title = 'Introduction'
    current_lines: list[str] = []
    section_index = 0

    def iter_block_items() -> list[Paragraph | Table]:
        body = document.element.body
        items: list[Paragraph | Table] = []
        for child in body.iterchildren():
            if child.tag.endswith('}p'):
                items.append(Paragraph(child, document))
                continue
            if child.tag.endswith('}tbl'):
                items.append(Table(child, document))
        return items

    def table_to_html(table: Table) -> str:
        rows_html: list[str] = []
        for row in table.rows:
            cell_html = ''.join(f'<td>{clean_text(cell.text) or ""}</td>' for cell in row.cells)
            rows_html.append(f'<tr>{cell_html}</tr>')
        return '<table>' + ''.join(rows_html) + '</table>'

    def flush_current_section() -> None:
        nonlocal section_index
        content = clean_text('\n'.join(current_lines))
        if not content:
            return
        sections.append(
            LoadedSection(
                text=content,
                metadata={
                    'section_type': 'docx_heading_block',
                    'section_title': current_title,
                    'section_index': section_index,
                },
            )
        )
        section_index += 1

    for block in iter_block_items():
        if isinstance(block, Paragraph):
            paragraph_text = block.text.strip()
            if not paragraph_text:
                continue

            style_name = getattr(block.style, 'name', '') or ''
            if style_name.startswith('Heading'):
                flush_current_section()
                current_title = paragraph_text
                current_lines = [paragraph_text]
                continue

            current_lines.append(paragraph_text)
            continue

        if isinstance(block, Table):
            current_lines.append(table_to_html(block))

    flush_current_section()
    if not sections:
        plain_text = clean_text('\n'.join(paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()))
        sections = [LoadedSection(text=plain_text, metadata={'section_type': 'full_text', 'section_index': 0})]

    loaded_document = LoadedDocument(
        filename=file_path.name,
        file_type='docx',
        parser_name='docx_loader',
        sections=sections,
    )
    logger.info(
        '[PARSER] selected: file=%s parser=NATIVE_%s sections=%s file_type=%s',
        file_path.name,
        loaded_document.parser_name.upper(),
        len(loaded_document.sections),
        loaded_document.file_type,
    )
    return loaded_document


def _load_doc_file(file_path: Path) -> LoadedDocument:
    """通过 OCR 解析旧版 `.doc` 文档。

    当前项目没有引入稳定的 `.doc` 原生解析器，因此这里直接把它归入 OCR 路线。
    """

    ocr_loaded_document = _try_load_with_ocr(file_path, file_type='doc')
    if ocr_loaded_document is not None:
        return ocr_loaded_document
    raise ValueError('Legacy .doc files require OCR service configuration to be enabled')


def build_loaded_document_from_text(filename: str, text: str) -> LoadedDocument:
    """为纯文本直传场景构造统一文档结构。"""

    suffix = Path(filename).suffix.lower().lstrip('.') or 'txt'
    section_type = 'markdown_heading' if suffix == 'md' else 'inline_text'
    sections = _split_markdown_sections(text) if suffix == 'md' else [
        LoadedSection(
            text=clean_text(text),
            metadata={
                'section_type': section_type,
                'section_index': 0,
            },
        )
    ]
    loaded_document = LoadedDocument(
        filename=filename,
        file_type=suffix,
        parser_name='inline_text_loader',
        sections=sections,
    )
    logger.info(
        '[PARSER] selected: file=%s parser=%s sections=%s file_type=%s',
        filename,
        loaded_document.parser_name,
        len(loaded_document.sections),
        loaded_document.file_type,
    )
    return loaded_document


LOADER_REGISTRY: dict[str, Callable[[Path], LoadedDocument]] = {
    'doc': _load_doc_file,
    'txt': _load_text_file,
    'md': _load_markdown_file,
    'pdf': _load_pdf_file,
    'docx': _load_docx_file,
}


def load_document(file_path: str | Path) -> LoadedDocument:
    """按文件扩展名分发到对应加载器。"""

    path = Path(file_path)
    file_type = path.suffix.lower().lstrip('.')
    loader = LOADER_REGISTRY.get(file_type)
    if loader is None:
        raise ValueError(f'Unsupported file type: {file_type or "unknown"}')
    return loader(path)
