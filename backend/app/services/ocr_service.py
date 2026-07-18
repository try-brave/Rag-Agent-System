from __future__ import annotations

import base64
import json
import logging
import time
from dataclasses import dataclass, field
from functools import lru_cache
from pathlib import Path
from typing import Any
from urllib.parse import urlencode, urlparse, parse_qsl, urlunparse
from urllib.request import Request, urlopen, build_opener, ProxyHandler

from docx import Document as DocxDocument
from pypdf import PdfReader

from app.config import get_settings
from app.utils.text import clean_text

logger = logging.getLogger(__name__)


@dataclass(slots=True)
class OCRDetectionDecision:
    """OCR 决策结果。

    之所以把“是否走 OCR”也抽象成一个对象，而不是只返回布尔值，是为了：
    - 把判断原因落日志；
    - 后续把原因透传到文档 summary / metadata；
    - 便于前端调试“为什么这次走了 OCR”。
    """

    should_use_ocr: bool
    file_type: str
    reasons: list[str] = field(default_factory=list)
    extracted_text_chars: int = 0
    image_count: int = 0
    table_count: int = 0
    empty_block_count: int = 0
    total_block_count: int = 0

    def to_metadata(self) -> dict[str, object]:
        """把决策结果压平成可落 metadata 的结构。"""

        return {
            'ocr_enabled': self.should_use_ocr,
            'ocr_reasons': self.reasons,
            'ocr_detected_text_chars': self.extracted_text_chars,
            'ocr_detected_image_count': self.image_count,
            'ocr_detected_table_count': self.table_count,
            'ocr_detected_empty_block_count': self.empty_block_count,
            'ocr_detected_total_block_count': self.total_block_count,
        }


class OCRService:
    """文档 OCR 服务。

    当前职责分两部分：
    1. 先对 PDF / 文档做启发式判断，判断是否“值得”走 OCR；
    2. 如果值得，再调用外部文档解析服务，把结果转换成 Markdown 文本。

    设计原则：
    - OCR 是增强能力，不是强依赖；
    - 判断命中时优先走 OCR；
    - OCR 调用失败时，调用方应可回退到原生解析。
    """

    SUPPORTED_FILE_TYPES = {'pdf', 'doc', 'docx'}

    def __init__(self) -> None:
        self.settings = get_settings()

    def can_use_ocr(self) -> bool:
        """判断当前环境是否具备 OCR 调用条件。"""

        current_settings = get_settings()
        return bool(
            current_settings.ocr_enabled
            and current_settings.ocr_task_url
            and current_settings.ocr_query_url
            and self._has_auth_material()
        )

    def _has_auth_material(self) -> bool:
        """判断是否具备 OCR 鉴权材料。"""

        # 获取最新的 settings，而不是初始化时的 self.settings
        current_settings = get_settings()
        return bool(
            current_settings.ocr_access_token
            or (current_settings.ocr_client_id and current_settings.ocr_client_secret)
        )

    def analyze_document(self, file_path: str | Path) -> OCRDetectionDecision:
        """根据文件内容判断是否建议走 OCR。"""

        path = Path(file_path)
        file_type = path.suffix.lower().lstrip('.')
        # 获取最新的 settings
        current_settings = get_settings()

        if not self.can_use_ocr() or file_type not in self.SUPPORTED_FILE_TYPES:
            return OCRDetectionDecision(should_use_ocr=False, file_type=file_type)

        if file_type == 'pdf':
            return self._analyze_pdf(path)
        if file_type == 'docx':
            return self._analyze_docx(path)

        # `.doc` 无法稳定做本地原生结构解析，若 OCR 已配置则直接走 OCR。
        return OCRDetectionDecision(
            should_use_ocr=True,
            file_type=file_type,
            reasons=['legacy_word_requires_ocr'],
        )

    def parse_to_markdown(self, file_path: str | Path) -> str:
        """调用 OCR 文档解析接口，并返回 Markdown 文本。"""

        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f'File not found: {path}')
        if not self.can_use_ocr():
            raise RuntimeError('OCR service is not configured')

        current_settings = get_settings()
        task_response = self._create_task(path)
        task_id = str(task_response['result']['task_id'])
        markdown_url = self._poll_markdown_url(task_id)
        markdown_content = self._download_text(markdown_url)
        cleaned_markdown = clean_text(markdown_content)
        logger.info('[OCR] markdown downloaded: file=%s task_id=%s markdown_chars=%s', path.name, task_id, len(cleaned_markdown))
        return cleaned_markdown

    def _analyze_pdf(self, file_path: Path) -> OCRDetectionDecision:
        """针对 PDF 做启发式判断。

        当前主要判断几种高风险场景：
        - 原生抽取文本极少，疑似扫描件；
        - 页面含图片且文本稀少，疑似图片型 PDF；
        - 行文本呈现明显表格信号，适合 OCR 保留结构。
        """

        reader = PdfReader(str(file_path))
        total_pages = len(reader.pages)
        extracted_text_chars = 0
        empty_page_count = 0
        image_count = 0
        table_like_line_count = 0

        for page in reader.pages:
            page_text = clean_text(page.extract_text() or '')
            extracted_text_chars += len(page_text)
            if len(page_text) < self.settings.ocr_pdf_min_page_chars:
                empty_page_count += 1

            table_like_line_count += self._count_table_like_lines(page_text)

            try:
                image_count += len(list(page.images))
            except Exception:  # noqa: BLE001
                # `pypdf` 对不同版本 PDF 的图片抽取兼容性不完全一致，这里只做 best effort。
                pass

        avg_text_chars = extracted_text_chars / max(total_pages, 1)
        reasons: list[str] = []
        if total_pages > 0 and empty_page_count == total_pages:
            reasons.append('all_pages_have_too_little_text')
        if total_pages > 0 and (empty_page_count / total_pages) >= self.settings.ocr_pdf_empty_page_ratio:
            reasons.append('most_pages_have_too_little_text')
        if image_count > 0 and avg_text_chars <= self.settings.ocr_pdf_low_text_avg_chars:
            reasons.append('pdf_contains_images_and_low_text_density')
        if table_like_line_count >= self.settings.ocr_table_like_line_threshold:
            reasons.append('pdf_contains_table_like_layout')

        decision = OCRDetectionDecision(
            should_use_ocr=bool(reasons),
            file_type='pdf',
            reasons=reasons,
            extracted_text_chars=extracted_text_chars,
            image_count=image_count,
            table_count=1 if table_like_line_count >= self.settings.ocr_table_like_line_threshold else 0,
            empty_block_count=empty_page_count,
            total_block_count=total_pages,
        )
        return decision

    def _analyze_docx(self, file_path: Path) -> OCRDetectionDecision:
        """针对 Docx 做启发式判断。

        Word 文档的正文通常能被原生抽取，但如果包含较多表格和内嵌图片，
        OCR 版本返回的 Markdown 更适合后续保留版式信息。
        """

        document = DocxDocument(str(file_path))
        paragraph_texts = [clean_text(paragraph.text) for paragraph in document.paragraphs if paragraph.text.strip()]
        extracted_text_chars = sum(len(text) for text in paragraph_texts)
        table_count = len(document.tables)
        image_count = len(document.inline_shapes)

        reasons: list[str] = []
        has_low_text_density = extracted_text_chars <= self.settings.ocr_docx_min_chars
        if image_count > 0:
            reasons.append('docx_contains_images')
        if table_count > 0:
            reasons.append('docx_contains_tables')
        if has_low_text_density:
            reasons.append('docx_text_density_is_low')

        decision = OCRDetectionDecision(
            should_use_ocr=bool(reasons),
            file_type='docx',
            reasons=reasons,
            extracted_text_chars=extracted_text_chars,
            image_count=image_count,
            table_count=table_count,
            empty_block_count=0,
            total_block_count=max(1, len(paragraph_texts)),
        )
        return decision

    @staticmethod
    def _count_table_like_lines(text: str) -> int:
        """粗略判断文本中是否带有表格布局信号。"""

        count = 0
        for line in text.splitlines():
            stripped = line.strip()
            if not stripped:
                continue
            has_delimiter = '|' in stripped or '\t' in stripped
            has_aligned_spaces = '  ' in stripped and len(stripped.split()) >= 3
            if has_delimiter or has_aligned_spaces:
                count += 1
        return count

    def _create_task(self, file_path: Path) -> dict[str, Any]:
        """创建 OCR 解析任务。"""

        with file_path.open('rb') as handle:
            encoded_file = base64.b64encode(handle.read()).decode('utf-8')

        payload = {
            'file_data': encoded_file,
            'file_url': '',
            'file_name': file_path.name,
        }
        current_settings = get_settings()
        response = self._post_form(self._append_access_token(current_settings.ocr_task_url), payload)
        logger.info('[OCR] task created: file=%s response_keys=%s', file_path.name, list(response.keys()))
        return response

    def _poll_markdown_url(self, task_id: str) -> str:
        """轮询 OCR 任务，直到拿到 Markdown 下载链接。"""

        current_settings = get_settings()
        for _ in range(current_settings.ocr_poll_max_attempts):
            query_response = self._post_form(
                self._append_access_token(current_settings.ocr_query_url),
                {'task_id': task_id},
            )
            result = query_response.get('result') or {}
            status = str(result.get('status') or '').lower()
            logger.info('[OCR] task polling: task_id=%s status=%s', task_id, status or 'unknown')

            if status == 'success':
                markdown_url = str(result.get('markdown_url') or '')
                if not markdown_url:
                    raise RuntimeError(f'OCR task succeeded but markdown_url is empty: task_id={task_id}')
                return markdown_url

            if status == 'failed':
                raise RuntimeError(f'OCR task failed: task_id={task_id}, error={result.get("task_error")}')

            time.sleep(current_settings.ocr_poll_interval_sec)

        raise RuntimeError(f'OCR task timed out: task_id={task_id}')

    @staticmethod
    def _download_text(url: str) -> str:
        """下载 OCR 返回的 Markdown 文本。"""

        request = Request(url=url, method='GET')
        opener = build_opener(ProxyHandler({}))
        with opener.open(request, timeout=30) as response:
            charset = response.headers.get_content_charset() or 'utf-8'
            return response.read().decode(charset, errors='replace')

    @staticmethod
    def _post_form(url: str, data: dict[str, Any]) -> dict[str, Any]:
        """发送 `application/x-www-form-urlencoded` POST 请求。"""

        encoded_body = urlencode({key: str(value) for key, value in data.items()}).encode('utf-8')
        request = Request(
            url=url,
            data=encoded_body,
            headers={'Content-Type': 'application/x-www-form-urlencoded'},
            method='POST',
        )
        opener = build_opener(ProxyHandler({}))
        with opener.open(request, timeout=60) as response:
            charset = response.headers.get_content_charset() or 'utf-8'
            raw_text = response.read().decode(charset, errors='replace')
        return json.loads(raw_text)

    @staticmethod
    def _append_access_token(url: str) -> str:
        """把 OCR access token 拼到请求 URL。"""

        access_token = _resolve_ocr_access_token()
        if not access_token:
            return url

        # 为了避免 urlunparse 把 url 结构弄乱，导致百度无法识别 token，
        # 直接按照测试脚本里的方式强拼接。
        separator = '&' if '?' in url else '?'
        return f"{url}{separator}access_token={access_token}"


def _resolve_ocr_access_token() -> str | None:
    """解析 OCR access token。

    每次调用都重新获取最新的配置单例，避免因为模块加载过早导致 .env 改动没生效。
    如果 access_token 为空且 client 凭证存在，则自动换取 token。
    """

    settings = get_settings()
    if settings.ocr_access_token:
        return settings.ocr_access_token

    if not settings.ocr_client_id or not settings.ocr_client_secret:
        return None

    token_url = (
        f'{settings.ocr_token_url}?grant_type=client_credentials'
        f'&client_id={settings.ocr_client_id}'
        f'&client_secret={settings.ocr_client_secret}'
    )
    request = Request(
        url=token_url,
        data=''.encode('utf-8'),
        headers={
            'Content-Type': 'application/json',
            'Accept': 'application/json',
        },
        method='POST',
    )
    opener = build_opener(ProxyHandler({}))
    with opener.open(request, timeout=30) as response:
        charset = response.headers.get_content_charset() or 'utf-8'
        payload = json.loads(response.read().decode(charset, errors='replace'))

    access_token = str(payload.get('access_token') or '').strip()
    if not access_token:
        raise RuntimeError(f'Failed to obtain OCR access token: response={payload}')

    logger.info('[OCR] access token resolved via client credentials')
    return access_token
